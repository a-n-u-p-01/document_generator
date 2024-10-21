from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
import re
from database import Template

app = Flask(__name__)
CORS(app)
@app.route('/upload', methods=['POST'])
def upload_file():
    print("Uploaded file received:", request.files)
    
    # the uploaded file is sent with key 'file' from client frontend
    doc_file = request.files['docFile']
    name = request.form.get('name')
    save = request.form.get('save')
    heading = request.form.get('heading')

    # Read the file content as binary before saving it
    doc_file_content = doc_file.read()
    file_path = 'uploaded.docx'
    
    with open(file_path, 'wb') as f:
        f.write(doc_file_content)

    # Create an instance of the Template class for data storage
    template_obj = Template() 

    # Save the template in the database with the binary content of the DOC file
    if save == '1': #true if doc need to store in database
        pdf_file_content = request.files["pdfFile"].read()
        print("File uploaded and saved.", template_obj.save_template(name, heading, pdf_file_content, doc_file_content))

    template_obj.close()

    # Load the document and find placeholders
    doc = Document(file_path)
    placeholders = []

    for p in doc.paragraphs:
        if '{{' in p.text:
            found_placeholders = re.findall(r'\{\{\s*(.*?)\s*\}\}', p.text)
            for placeholder in found_placeholders:
                if placeholder not in placeholders:  # Check if already added
                    placeholders.append(placeholder)

    print(f"Found placeholders: {placeholders}")

    return {'placeholders': placeholders}

@app.route('/replace', methods=['POST'])
def replace_placeholders():
    data = request.json

    if not data:
        return {"error": "No data provided"}, 400

    print(f"Received data for replacement: {data}")

    try:
        doc = Document('uploaded.docx')
        updated = False
        placeholder_presents = [f"{{{key}}}" for key in data.keys()]
        print(placeholder_presents)
        for index,p in enumerate(doc.paragraphs):
            for index,run in enumerate(p.runs):
                print(f"run:{index} text: {run.text}")
                matches = re.findall(r'\{\{(.*?)\}\}', run.text)
                
                for match in matches:
                    if match in data:
                        value = data[match]
                        parts = re.split(r'\{\{' + re.escape(match) + r'\}\}', run.text)
                        run.text = parts[0]+value+parts[1]
                        updated=True
                      
                        if f"{{{match}}}" in placeholder_presents:
                            placeholder_presents.remove(f"{{{match}}}")
                            print("Remaining Place Holder+++",placeholder_presents)
                            print(f"Removed {match} from placeholders")
                        
        if len(placeholder_presents)>0:
            print("Remaining Place Holder",placeholder_presents)
            return {"placeholders":placeholder_presents},403

        if updated:
            doc.save('uploaded.docx')
            print("Document updated and saved as uploaded.docx.")
            return send_file('uploaded.docx', as_attachment=True)
        else:
            print("No placeholders found to replace.")
            return {"error": "No placeholders found"}, 400

    except Exception as e:
        print(f"Error updating document: {e}")
        return {"error": "Error updating document"}, 500

#fetch all the templates from data
@app.route('/templates', methods=['GET'])
def get_templates():
    # Create an instance of the Template class
    template_obj = Template()
    templates = template_obj.getAll()  # Get all templates
    template_obj.close()
    
    return jsonify(templates)  # Return templates as JSON

if __name__ == '__main__':
    app.run(debug=True)